"""
gen_word_nacional_py.py — Pure Python Nacional Document Generator (python-docx)

Rewrite of gen_word_nacional.js using python-docx only.
No Node.js dependency. Professional styling maintained with XML manipulation.

Usage:
    from gen_word_nacional_py import generate_nacional_docx
    docx_bytes = generate_nacional_docx(datos_dict)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from io import BytesIO


# ═══ Colors ═══
C = {
    "BLUE": "2F5496",
    "GRAY": "666666",
    "WHITE": "FFFFFF",
    "HEADER_BG": "2F5496",
    "ALT_ROW": "D6E4F0",
    "BLACK": "000000",
}


def fmt_num(val, dec=2):
    """Format number with thousands separator (es-PE locale)."""
    if val is None or val == "":
        return "0"
    try:
        n = float(val)
        if dec == 0:
            return f"{n:,.0f}".replace(",", " ").replace(".", ",").replace(" ", ".")
        return f"{n:,.{dec}f}".replace(",", " ").replace(".", ",").replace(" ", ".")
    except (ValueError, TypeError):
        return "0"


def set_cell_background(cell, fill_color):
    """Set cell background color using XML shading."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Set cell borders. Usage: set_cell_border(cell, top={...}, bottom={...})"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(r'<w:tcBorders %s><w:top w:val="single" w:sz="4" w:space="0" w:color="BBBBBB"/>'
                          r'<w:left w:val="single" w:sz="4" w:space="0" w:color="BBBBBB"/>'
                          r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="BBBBBB"/>'
                          r'<w:right w:val="single" w:sz="4" w:space="0" w:color="BBBBBB"/></w:tcBorders>' % nsdecls("w"))
    tcPr.append(tcBorders)


def set_paragraph_bottom_border(paragraph, color="C0392B", size="6"):
    """Add bottom border to paragraph."""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="{size}" w:space="1" w:color="{color}"/></w:pBdr>'
    )
    pPr.append(pBdr)


def create_table(doc, headers, rows, col_widths_twips=None):
    """Create a professional table with headers and data rows."""
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Light Grid Accent 1"

    # Auto width if not specified
    if col_widths_twips is None:
        col_widths_twips = [9360 // len(headers)] * len(headers)

    # Set column widths
    for i, width in enumerate(col_widths_twips):
        table.columns[i].width = width

    # Header row
    header_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        cell = header_cells[i]
        cell.text = header_text
        set_cell_background(cell, C["HEADER_BG"])

        # Format header text
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(9)
        
        # Add proper borders
        set_cell_border(cell)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        is_total = row_data[0].upper() == "TOTAL" if row_data else False
        is_alt = row_idx % 2 == 0
        bg_color = C["HEADER_BG"] if is_total else (C["ALT_ROW"] if is_alt else None)
        text_color = C["WHITE"] if is_total else C["BLACK"]

        table_row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            cell = table_row.cells[col_idx]
            cell.text = str(cell_text) if cell_text else ""
            
            if bg_color:
                set_cell_background(cell, bg_color)
            
            # Format cell text
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if col_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            
            for run in paragraph.runs:
                run.font.bold = is_total
                run.font.size = Pt(9)
                if text_color == C["WHITE"]:
                    run.font.color.rgb = RGBColor(255, 255, 255)
            
            # Add borders
            set_cell_border(cell)

    return table


def generate_nacional_docx(datos):
    """
    Generate NACIONAL ayuda memoria document.
    
    Args:
        datos (dict): Data dictionary with keys:
            - fecha_corte
            - total_avisos, total_ajustados, pct_ajustados
            - monto_indemnizado, monto_desembolsado, productores_desembolso
            - prima_total, prima_neta, sup_asegurada, prod_asegurados
            - indice_siniestralidad, pct_desembolso, deptos_con_desembolso
            - empresas_text
            - cuadro1, cuadro2, cuadro3 (list of dicts)
            - total_lluvia, pct_lluvia, lluvia_desc
            - top3_lluvia_text, top3_siniestros_text
    
    Returns:
        bytes: DOCX file content
    """
    doc = Document()
    
    # Set page margins: 1200 twips = 0.833 inches
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.833)
        section.bottom_margin = Inches(0.833)
        section.left_margin = Inches(0.833)
        section.right_margin = Inches(0.833)
    
    # ═══ TITLE ═══
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(12)
    
    # "AYUDA MEMORIA: " (14pt)
    run = title_para.add_run("AYUDA MEMORIA: ")
    run.font.name = "Arial"
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(47, 84, 150)  # #2F5496
    
    # "RESUMEN OPERATIVIDAD SAC 2025-2026" (12pt)
    run = title_para.add_run("RESUMEN OPERATIVIDAD SAC 2025-2026")
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(47, 84, 150)
    
    # Subtitle: "(al fecha)"
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_para.paragraph_format.space_after = Pt(18)
    run = subtitle_para.add_run(f"(al {datos['fecha_corte']})")
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(102, 102, 102)  # #666666
    
    # ═══ ACTIVATION SECTION ═══
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(8)
    run = h1.add_run("Activación del SAC")
    run.font.name = "Arial"
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(47, 84, 150)
    
    # Body text
    body = doc.add_paragraph("El Seguro Agrícola Catastrófico se activa mediante el siguiente procedimiento:")
    body.paragraph_format.space_after = Pt(6)
    for run in body.runs:
        run.font.name = "Arial"
        run.font.size = Pt(11)
    
    # Steps (numbered list)
    pasos = [
        "El productor reporta el siniestro a la Agencia/Oficina Agraria (presencial o telefónicamente).",
        "La Agencia Agraria, vía la DRA, notifica a la aseguradora por correo electrónico dentro de 7 días calendario.",
        "La aseguradora designa un perito ajustador para evaluar daños en campo dentro de 15 días calendario.",
        "El perito evalúa los daños con un agente agrario y elabora el acta de ajuste.",
        "Si se confirma pérdida catastrófica, se coordina el empadronamiento de agricultores dentro de 20 días calendario.",
        "La aseguradora abre cuentas bancarias y paga S/ 1,000 por hectárea asegurada dentro de 15 días hábiles tras la aprobación del padrón.",
    ]
    
    for paso in pasos:
        p = doc.add_paragraph(paso, style="List Number")
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(11)
            run.font.bold = True
    
    doc.add_paragraph()  # spacing
    
    # ═══ GENERAL DATA SECTION ═══
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(8)
    run = h1.add_run("Datos Generales a Nivel Nacional")
    run.font.name = "Arial"
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(47, 84, 150)
    
    # Bullets
    bullets = [
        f"Empresas aseguradoras: {datos['empresas_text']}.",
        f"Prima total (con IGV): S/ {fmt_num(datos['prima_total'])} | Prima neta (sin IGV): S/ {fmt_num(datos['prima_neta'])}",
        f"Superficie asegurada: {fmt_num(datos['sup_asegurada'])} hectáreas en 24 departamentos.",
        f"Productores asegurados (estimados): {fmt_num(datos['prod_asegurados'], 0)} / Suma asegurada por hectárea: S/ 1,000.00",
        f"Avisos de siniestros: {fmt_num(datos['total_avisos'], 0)} reportados | {fmt_num(datos['total_ajustados'], 0)} ajustados ({datos['pct_ajustados']}%)",
        f"Indemnizaciones reconocidas: S/ {fmt_num(datos['monto_indemnizado'])} | Índice de siniestralidad: {datos['indice_siniestralidad']}%",
        f"Desembolsos realizados: S/ {fmt_num(datos['monto_desembolsado'])} ({datos['pct_desembolso']}%) a {fmt_num(datos['productores_desembolso'], 0)} productores en {datos['deptos_con_desembolso']} de 24 departamentos.",
    ]
    
    for bullet_text in bullets:
        p = doc.add_paragraph(bullet_text, style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(11)
            run.font.bold = True
    
    doc.add_paragraph()  # spacing
    
    # ═══ TABLA 1 ═══
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(8)
    run = h1.add_run("Cuadro 1: Primas y Cobertura por Departamento")
    run.font.name = "Arial"
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(47, 84, 150)
    
    if datos.get("cuadro1") and len(datos["cuadro1"]) > 0:
        headers = ["Departamento", "Prima Total (S/)", "Hectáreas Aseguradas", "Suma Asegurada Máxima (S/)"]
        rows = []
        
        for item in datos["cuadro1"]:
            rows.append([
                item.get("departamento", ""),
                fmt_num(item.get("prima_total", 0)),
                fmt_num(item.get("hectareas", 0)),
                fmt_num(item.get("suma_asegurada", 0)),
            ])
        
        # Add TOTAL row
        if len(rows) > 0:
            total_prima = sum(float(item.get("prima_total", 0) or 0) for item in datos["cuadro1"])
            total_ha = sum(float(item.get("hectareas", 0) or 0) for item in datos["cuadro1"])
            total_suma = sum(float(item.get("suma_asegurada", 0) or 0) for item in datos["cuadro1"])
            rows.append(["TOTAL", fmt_num(total_prima), fmt_num(total_ha), fmt_num(total_suma)])
        
        col_widths = [1920, 1760, 1760, 2048]
        create_table(doc, headers, rows, col_widths)
    
    doc.add_paragraph()  # spacing
    
    # ═══ TABLA 2 ═══
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(8)
    run = h1.add_run("Cuadro 2: Indemnizaciones y Desembolsos por Departamento")
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(47, 84, 150)
    
    subtitle = h1.insert_paragraph_before(f"(al {datos['fecha_corte']})")
    subtitle.paragraph_format.space_after = Pt(6)
    for run in subtitle.runs:
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(102, 102, 102)
    
    if datos.get("cuadro2") and len(datos["cuadro2"]) > 0:
        headers = ["Departamento", "Ha Indemnizadas", "Monto Indemnizado (S/)", "Monto Desembolsado (S/)", "Productores con Desembolso"]
        rows = []
        
        for item in datos["cuadro2"]:
            rows.append([
                item.get("departamento", ""),
                fmt_num(item.get("ha_indemnizadas", 0)),
                fmt_num(item.get("monto_indemnizado", 0)),
                fmt_num(item.get("monto_desembolsado", 0)),
                fmt_num(item.get("productores", 0), 0),
            ])
        
        # Add TOTAL row
        if len(rows) > 0:
            total_ha_i = sum(float(item.get("ha_indemnizadas", 0) or 0) for item in datos["cuadro2"])
            total_monto_i = sum(float(item.get("monto_indemnizado", 0) or 0) for item in datos["cuadro2"])
            total_monto_d = sum(float(item.get("monto_desembolsado", 0) or 0) for item in datos["cuadro2"])
            total_prod = sum(float(item.get("productores", 0) or 0) for item in datos["cuadro2"])
            rows.append(["TOTAL", fmt_num(total_ha_i), fmt_num(total_monto_i), fmt_num(total_monto_d), fmt_num(total_prod, 0)])
        
        col_widths = [1600, 1280, 1680, 1680, 1248]
        create_table(doc, headers, rows, col_widths)
    
    doc.add_paragraph()  # spacing
    
    # ═══ TABLA 3 ═══
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(8)
    run = h1.add_run("Cuadro 3: Eventos Asociados a Lluvias Intensas por Departamento")
    run.font.name = "Arial"
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(47, 84, 150)
    
    # Descriptive paragraph
    desc_para = doc.add_paragraph()
    desc_para.paragraph_format.space_after = Pt(8)
    run = desc_para.add_run(
        f"Se registran {fmt_num(datos['total_lluvia'], 0)} avisos por eventos asociados a lluvias intensas "
        f"({datos['pct_lluvia']}% del total), que incluyen {datos['lluvia_desc']}. "
        f"Los departamentos más afectados son {datos['top3_lluvia_text']}."
    )
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.bold = True
    
    if datos.get("cuadro3") and len(datos["cuadro3"]) > 0:
        headers = ["Departamento", "Avisos", "Ha Indemn.", "Monto Indemnizado (S/)", "Monto Desembolsado (S/)", "Productores"]
        rows = []
        
        for item in datos["cuadro3"]:
            rows.append([
                item.get("departamento", ""),
                fmt_num(item.get("avisos", 0), 0),
                fmt_num(item.get("ha_indemn", 0)),
                fmt_num(item.get("monto_indemnizado", 0)),
                fmt_num(item.get("monto_desembolsado", 0)),
                fmt_num(item.get("productores", 0), 0),
            ])
        
        col_widths = [1440, 800, 1040, 1680, 1680, 848]
        create_table(doc, headers, rows, col_widths)
    
    doc.add_paragraph()  # spacing
    
    # ═══ FINAL NOTE ═══
    note_para = doc.add_paragraph()
    note_para.paragraph_format.space_after = Pt(4)
    
    run = note_para.add_run("Nota: ")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.italic = True
    
    run = note_para.add_run("La vigencia de la póliza es del 01/08/2025 al 01/08/2026. ")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.italic = True
    
    run = note_para.add_run(datos.get("top3_siniestros_text", ""))
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.italic = True
    
    # ═══ RETURN BYTES ═══
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


if __name__ == "__main__":
    # Quick test
    import json
    
    sample_data = {
        "fecha_corte": "27/02/2026",
        "total_avisos": 1250,
        "total_ajustados": 1100,
        "pct_ajustados": "88.0%",
        "monto_indemnizado": 5500000.00,
        "monto_desembolsado": 4200000.00,
        "productores_desembolso": 890,
        "prima_total": 12500000.00,
        "prima_neta": 11200000.00,
        "sup_asegurada": 125000.00,
        "prod_asegurados": 8900,
        "indice_siniestralidad": 49.0,
        "pct_desembolso": "76.4%",
        "deptos_con_desembolso": 20,
        "empresas_text": "La Positiva y Rímac",
        "cuadro1": [
            {"departamento": "Lambayeque", "prima_total": 1200000, "hectareas": 15000, "suma_asegurada": 15000000},
            {"departamento": "Piura", "prima_total": 1100000, "hectareas": 14000, "suma_asegurada": 14000000},
            {"departamento": "Áncash", "prima_total": 950000, "hectareas": 12000, "suma_asegurada": 12000000},
        ],
        "cuadro2": [
            {"departamento": "Lambayeque", "ha_indemnizadas": 8900, "monto_indemnizado": 2500000, "monto_desembolsado": 1900000, "productores": 450},
            {"departamento": "Piura", "ha_indemnizadas": 7200, "monto_indemnizado": 1800000, "monto_desembolsado": 1400000, "productores": 320},
        ],
        "cuadro3": [
            {"departamento": "Lambayeque", "avisos": 450, "ha_indemn": 5600, "monto_indemnizado": 2800000, "monto_desembolsado": 2100000, "productores": 280},
        ],
        "total_lluvia": 980,
        "pct_lluvia": "78.4%",
        "lluvia_desc": "inundación, huayco, lluvias excesivas y deslizamiento",
        "top3_lluvia_text": "Lambayeque (450 avisos), Piura (320 avisos), Áncash (210 avisos)",
        "top3_siniestros_text": "Los siniestros principales son inundación (65%), lluvias excesivas (20%) y huayco (15%).",
    }
    
    docx_bytes = generate_nacional_docx(sample_data)
    print(f"Generated NACIONAL document: {len(docx_bytes)} bytes")
    print("✓ No import errors, document structure valid")
