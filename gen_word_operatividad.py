"""
gen_word_operatividad.py — Generador de Ayuda Memoria OPERATIVIDAD SAC
======================================================================
Genera un documento Word con el formato de "AM - OPERATIVIDAD SAC"
que incluye:
  - Avisos de siniestros por empresa y departamento
  - Resultados: ajustes, indemnizaciones, siniestralidad
  - Tabla de siniestralidad por empresa/departamento
  - Tabla de coberturas (complementaria vs catastrófica)
  - Tabla de cultivos priorizados vs no priorizados
  - Tabla de desembolsos por empresa/departamento

Usa python-docx (misma tecnología que los demás generadores).
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from io import BytesIO
import pandas as pd
import numpy as np


# ═══ COLORES ═══
C = {
    "BLUE": "1F4E79",
    "LIGHT_BLUE": "2E75B6",
    "WHITE": "FFFFFF",
    "BLACK": "000000",
    "HEADER_BG": "1F4E79",
    "SUBTOTAL_BG": "D6E4F0",
    "ALT_ROW": "F2F7FB",
    "GRAY": "666666",
}


def fmt(val, dec=2):
    """Formatea número con separador de miles estilo S/."""
    if val is None or (isinstance(val, float) and np.isnan(val)):
        return "S/ -"
    try:
        n = float(val)
        if n == 0:
            return "S/ -"
        formatted = f"{n:,.{dec}f}"
        return f"S/ {formatted}"
    except (ValueError, TypeError):
        return "S/ -"


def fmt_n(val, dec=2):
    """Formatea número sin prefijo S/."""
    if val is None or (isinstance(val, float) and np.isnan(val)):
        return "-"
    try:
        n = float(val)
        if n == 0:
            return "-"
        return f"{n:,.{dec}f}"
    except (ValueError, TypeError):
        return "-"


def fmt_pct(val, dec=2):
    """Formatea porcentaje."""
    if val is None or (isinstance(val, float) and np.isnan(val)):
        return "0.00%"
    try:
        return f"{float(val):,.{dec}f}%"
    except (ValueError, TypeError):
        return "0.00%"


def fmt_int(val):
    """Formatea entero con separador de miles."""
    try:
        return f"{int(float(val)):,}"
    except (ValueError, TypeError):
        return "-"


# ═══════════════════════════════════════════════════════════════════
# UTILIDADES PARA TABLAS
# ═══════════════════════════════════════════════════════════════════

def _set_bg(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_borders(cell, color="BBBBBB"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)


def _set_cell_width(cell, width_twips):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{width_twips}" w:type="dxa"/>')
    tcPr.append(tcW)


def _write_cell(cell, text, bold=False, size=8, align=WD_ALIGN_PARAGRAPH.LEFT,
                font_color=None, bg_color=None):
    """Escribe en celda con formato."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(str(text))
    run.font.name = "Arial Narrow"
    run.font.size = Pt(size)
    run.font.bold = bold
    if font_color:
        run.font.color.rgb = RGBColor.from_string(font_color)
    if bg_color:
        _set_bg(cell, bg_color)
    _set_borders(cell)


def _merge_vertical(table, col, start_row, end_row):
    """Merge celdas verticalmente."""
    cell_start = table.cell(start_row, col)
    cell_end = table.cell(end_row, col)
    cell_start.merge(cell_end)


# ═══════════════════════════════════════════════════════════════════
# FUNCIÓN PARA AGREGAR TEXTO FORMATEADO
# ═══════════════════════════════════════════════════════════════════

def _add_heading(doc, text, level=1):
    """Agrega un encabezado con formato."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.bold = True
    run.underline = True
    if level == 1:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor.from_string(C["BLACK"])
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor.from_string(C["BLACK"])
    return p


def _add_bullet(doc, text, bold_prefix="", indent_level=0):
    """Agrega un bullet point."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    if indent_level > 0:
        p.paragraph_format.left_indent = Inches(0.5 * indent_level)

    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.font.bold = True

    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    return p


def _add_subbullet(doc, text):
    """Agrega un sub-bullet (nivel 2)."""
    p = doc.add_paragraph(style="List Bullet 2")
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    return p


def _add_body(doc, text, bold=False, italic=False, size=10):
    """Agrega un párrafo de texto normal."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    return p


# ═══════════════════════════════════════════════════════════════════
# PREPARACIÓN DE DATOS
# ═══════════════════════════════════════════════════════════════════

def _prepare_operatividad_data(datos):
    """
    Prepara todas las agregaciones necesarias para el documento
    a partir del dict datos de data_processor.
    """
    midagri = datos["midagri"]        # DataFrame combinado (LP + Rímac)
    materia = datos["materia"]        # Materia asegurada (estático)

    # ─── Mapeo departamento → empresa ───
    depto_empresa = {}
    if "EMPRESA_ASEGURADORA" in materia.columns and "DEPARTAMENTO" in materia.columns:
        for _, row in materia.iterrows():
            d = str(row["DEPARTAMENTO"]).strip().upper()
            e = str(row["EMPRESA_ASEGURADORA"]).strip().upper()
            depto_empresa[d] = e

    # Asignar empresa a cada registro
    if "DEPARTAMENTO" in midagri.columns:
        midagri = midagri.copy()
        midagri["EMPRESA"] = midagri["DEPARTAMENTO"].map(depto_empresa).fillna("OTROS")
    else:
        midagri["EMPRESA"] = "OTROS"

    # Normalizar nombres de empresa
    def _normalize_empresa(e):
        eu = str(e).upper()
        if "POSITIVA" in eu:
            return "LA POSITIVA"
        elif "RIMAC" in eu or "RÍMAC" in eu:
            return "RÍMAC"
        return eu
    midagri["EMPRESA"] = midagri["EMPRESA"].apply(_normalize_empresa)

    # ─── Prima neta por departamento (desde materia) ───
    prima_por_depto = {}
    if "PRIMA_NETA" in materia.columns and "DEPARTAMENTO" in materia.columns:
        for _, row in materia.iterrows():
            d = str(row["DEPARTAMENTO"]).strip().upper()
            prima_por_depto[d] = float(row.get("PRIMA_NETA", 0) or 0)

    fecha_corte = datos["fecha_corte"]

    # ═══ 1. AVISOS POR EMPRESA ═══
    avisos_by_empresa = midagri.groupby("EMPRESA").size().to_dict()
    total_avisos = len(midagri)

    avisos_lp = avisos_by_empresa.get("LA POSITIVA", 0)
    avisos_rimac = avisos_by_empresa.get("RÍMAC", 0)

    # ═══ 2. TOP DEPARTAMENTOS POR AVISOS ═══
    avisos_by_depto = midagri.groupby("DEPARTAMENTO").size().sort_values(ascending=False)
    top4_deptos = avisos_by_depto.head(4)
    top4_total = top4_deptos.sum()
    top4_pct = (top4_total / total_avisos * 100) if total_avisos > 0 else 0

    # ═══ 3. TOP TIPOS DE SINIESTRO ═══
    avisos_by_tipo = midagri["TIPO_SINIESTRO"].value_counts() if "TIPO_SINIESTRO" in midagri.columns else pd.Series()
    top3_tipos = avisos_by_tipo.head(3)
    top3_total = top3_tipos.sum()
    top3_pct = (top3_total / total_avisos * 100) if total_avisos > 0 else 0

    # ═══ 4. AJUSTADOS POR EMPRESA ═══
    if "ESTADO_INSPECCION" in midagri.columns:
        ajustados_mask = midagri["ESTADO_INSPECCION"].astype(str).str.upper() == "CERRADO"
    elif "ESTADO_SINIESTRO" in midagri.columns:
        ajustados_mask = midagri["ESTADO_SINIESTRO"].astype(str).str.upper() == "CONCRETADO"
    else:
        ajustados_mask = pd.Series([False] * len(midagri))

    total_ajustados = ajustados_mask.sum()
    pct_ajustados = (total_ajustados / total_avisos * 100) if total_avisos > 0 else 0

    ajust_lp = ajustados_mask[midagri["EMPRESA"] == "LA POSITIVA"].sum()
    ajust_rimac = ajustados_mask[midagri["EMPRESA"] == "RÍMAC"].sum()
    pct_ajust_lp = (ajust_lp / avisos_lp * 100) if avisos_lp > 0 else 0
    pct_ajust_rimac = (ajust_rimac / avisos_rimac * 100) if avisos_rimac > 0 else 0

    # ═══ 5. INDEMNIZACIÓN TOTAL ═══
    monto_indemnizado = midagri["INDEMNIZACION"].sum() if "INDEMNIZACION" in midagri.columns else 0

    # ═══ 6. TABLA SINIESTRALIDAD POR EMPRESA/DEPARTAMENTO ═══
    tabla_siniestralidad = []
    for empresa in ["LA POSITIVA", "RÍMAC"]:
        df_emp = midagri[midagri["EMPRESA"] == empresa]
        deptos = sorted(df_emp["DEPARTAMENTO"].unique())

        emp_indemn_total = 0
        emp_sup_total = 0
        emp_prima_total = 0

        filas_empresa = []
        for depto in deptos:
            df_d = df_emp[df_emp["DEPARTAMENTO"] == depto]
            indemn = df_d["INDEMNIZACION"].sum() if "INDEMNIZACION" in df_d.columns else 0
            sup_ind = df_d["SUP_INDEMNIZADA"].sum() if "SUP_INDEMNIZADA" in df_d.columns else 0
            prima = prima_por_depto.get(depto, 0)
            idx_sin = (indemn / prima * 100) if prima > 0 else 0

            emp_indemn_total += indemn
            emp_sup_total += sup_ind
            emp_prima_total += prima

            filas_empresa.append({
                "empresa": empresa,
                "departamento": depto.title(),
                "indemnizacion": indemn,
                "sup_indemnizada": sup_ind,
                "prima_neta": prima,
                "indice": idx_sin,
            })

        # Ordenar por índice de siniestralidad descendente
        filas_empresa.sort(key=lambda x: x["indice"], reverse=True)
        tabla_siniestralidad.extend(filas_empresa)

        # Subtotal empresa
        idx_emp = (emp_indemn_total / emp_prima_total * 100) if emp_prima_total > 0 else 0
        tabla_siniestralidad.append({
            "empresa": f"Total {empresa}",
            "departamento": "",
            "indemnizacion": emp_indemn_total,
            "sup_indemnizada": emp_sup_total,
            "prima_neta": emp_prima_total,
            "indice": idx_emp,
            "is_subtotal": True,
        })

    # Total general
    prima_total_neta = datos.get("prima_neta", 0)
    idx_general = (monto_indemnizado / prima_total_neta * 100) if prima_total_neta > 0 else 0
    sup_ind_total = midagri["SUP_INDEMNIZADA"].sum() if "SUP_INDEMNIZADA" in midagri.columns else 0
    tabla_siniestralidad.append({
        "empresa": "Total general",
        "departamento": "",
        "indemnizacion": monto_indemnizado,
        "sup_indemnizada": sup_ind_total,
        "prima_neta": prima_total_neta,
        "indice": idx_general,
        "is_total": True,
    })

    # Siniestralidad por empresa
    indemn_lp = midagri[midagri["EMPRESA"] == "LA POSITIVA"]["INDEMNIZACION"].sum() if "INDEMNIZACION" in midagri.columns else 0
    indemn_rimac = midagri[midagri["EMPRESA"] == "RÍMAC"]["INDEMNIZACION"].sum() if "INDEMNIZACION" in midagri.columns else 0
    prima_lp = sum(prima_por_depto.get(d, 0) for d in midagri[midagri["EMPRESA"] == "LA POSITIVA"]["DEPARTAMENTO"].unique())
    prima_rimac = sum(prima_por_depto.get(d, 0) for d in midagri[midagri["EMPRESA"] == "RÍMAC"]["DEPARTAMENTO"].unique())
    idx_lp = (indemn_lp / prima_lp * 100) if prima_lp > 0 else 0
    idx_rimac = (indemn_rimac / prima_rimac * 100) if prima_rimac > 0 else 0

    # ═══ 7. TABLA COBERTURAS (complementaria vs catastrófica) ═══
    tabla_coberturas = []
    if "TIPO_COBERTURA" in midagri.columns:
        for empresa in ["LA POSITIVA", "RÍMAC"]:
            df_emp = midagri[midagri["EMPRESA"] == empresa]
            deptos = sorted(df_emp["DEPARTAMENTO"].unique())

            emp_comp = 0
            emp_cat = 0
            emp_total = 0

            for depto in deptos:
                df_d = df_emp[df_emp["DEPARTAMENTO"] == depto]
                # Complementaria
                comp_mask = df_d["TIPO_COBERTURA"].astype(str).str.upper().str.contains("COMPLEMENT", na=False)
                cat_mask = df_d["TIPO_COBERTURA"].astype(str).str.upper().str.contains("CATASTR", na=False)

                val_comp = df_d.loc[comp_mask, "INDEMNIZACION"].sum() if "INDEMNIZACION" in df_d.columns else 0
                val_cat = df_d.loc[cat_mask, "INDEMNIZACION"].sum() if "INDEMNIZACION" in df_d.columns else 0
                val_total = df_d["INDEMNIZACION"].sum() if "INDEMNIZACION" in df_d.columns else 0

                emp_comp += val_comp
                emp_cat += val_cat
                emp_total += val_total

                if val_total > 0:
                    tabla_coberturas.append({
                        "empresa": empresa,
                        "departamento": depto.title(),
                        "complementaria": val_comp,
                        "catastrofica": val_cat,
                        "total": val_total,
                    })

            # Subtotal empresa
            tabla_coberturas.append({
                "empresa": f"Total {empresa}",
                "departamento": "",
                "complementaria": emp_comp,
                "catastrofica": emp_cat,
                "total": emp_total,
                "is_subtotal": True,
            })

        # Total general coberturas
        total_comp = sum(r["complementaria"] for r in tabla_coberturas if not r.get("is_subtotal"))
        total_cat = sum(r["catastrofica"] for r in tabla_coberturas if not r.get("is_subtotal"))
        tabla_coberturas.append({
            "empresa": "Total general",
            "departamento": "",
            "complementaria": total_comp,
            "catastrofica": total_cat,
            "total": total_comp + total_cat,
            "is_total": True,
        })

    # ═══ 8. TABLA CULTIVOS PRIORIZADOS ═══
    tabla_priorizados = []
    if "PRIORIZADO" in midagri.columns:
        for empresa in ["LA POSITIVA", "RÍMAC"]:
            df_emp = midagri[midagri["EMPRESA"] == empresa]

            for prio_val in ["PRIORIZADO", "NO PRIORIZADO"]:
                mask = df_emp["PRIORIZADO"].astype(str).str.upper().str.contains(
                    "NO" if prio_val == "NO PRIORIZADO" else "^(?!.*NO)", regex=True, na=False
                )
                if prio_val == "NO PRIORIZADO":
                    mask = df_emp["PRIORIZADO"].astype(str).str.upper().str.contains("NO", na=False)
                else:
                    mask = ~df_emp["PRIORIZADO"].astype(str).str.upper().str.contains("NO", na=False)

                sup = df_emp.loc[mask, "SUP_INDEMNIZADA"].sum() if "SUP_INDEMNIZADA" in df_emp.columns else 0
                ind = df_emp.loc[mask, "INDEMNIZACION"].sum() if "INDEMNIZACION" in df_emp.columns else 0

                tabla_priorizados.append({
                    "empresa": empresa,
                    "cultivo": prio_val,
                    "sup_indemnizada": sup,
                    "indemnizacion": ind,
                })

            # Subtotal
            emp_sup = df_emp["SUP_INDEMNIZADA"].sum() if "SUP_INDEMNIZADA" in df_emp.columns else 0
            emp_ind = df_emp["INDEMNIZACION"].sum() if "INDEMNIZACION" in df_emp.columns else 0
            tabla_priorizados.append({
                "empresa": f"Total {empresa}",
                "cultivo": "",
                "sup_indemnizada": emp_sup,
                "indemnizacion": emp_ind,
                "is_subtotal": True,
            })

        tabla_priorizados.append({
            "empresa": "Total general",
            "cultivo": "",
            "sup_indemnizada": sup_ind_total,
            "indemnizacion": monto_indemnizado,
            "is_total": True,
        })

    # ═══ 9. TABLA DESEMBOLSOS ═══
    tabla_desembolsos = []
    for empresa in ["LA POSITIVA", "RÍMAC"]:
        df_emp = midagri[midagri["EMPRESA"] == empresa]
        deptos = sorted(df_emp["DEPARTAMENTO"].unique())

        emp_indemn = 0
        emp_desemb = 0
        emp_prod = 0

        filas_emp = []
        for depto in deptos:
            df_d = df_emp[df_emp["DEPARTAMENTO"] == depto]
            indemn = df_d["INDEMNIZACION"].sum() if "INDEMNIZACION" in df_d.columns else 0
            desemb = df_d["MONTO_DESEMBOLSADO"].sum() if "MONTO_DESEMBOLSADO" in df_d.columns else 0
            prod = df_d["N_PRODUCTORES"].sum() if "N_PRODUCTORES" in df_d.columns else 0
            pct = (desemb / indemn * 100) if indemn > 0 else 0

            emp_indemn += indemn
            emp_desemb += desemb
            emp_prod += prod

            filas_emp.append({
                "empresa": empresa,
                "departamento": depto.title(),
                "indemnizacion": indemn,
                "desembolso": desemb,
                "pct_desembolso": pct,
                "productores": int(prod),
            })

        # Ordenar por % desembolso descendente
        filas_emp.sort(key=lambda x: x["pct_desembolso"], reverse=True)
        tabla_desembolsos.extend(filas_emp)

        pct_emp = (emp_desemb / emp_indemn * 100) if emp_indemn > 0 else 0
        tabla_desembolsos.append({
            "empresa": f"Total {empresa}",
            "departamento": "",
            "indemnizacion": emp_indemn,
            "desembolso": emp_desemb,
            "pct_desembolso": pct_emp,
            "productores": int(emp_prod),
            "is_subtotal": True,
        })

    monto_desembolsado = datos.get("monto_desembolsado", 0)
    productores = datos.get("productores_desembolso", 0)
    pct_desembolso = datos.get("pct_desembolso", 0)
    tabla_desembolsos.append({
        "empresa": "Total general",
        "departamento": "",
        "indemnizacion": monto_indemnizado,
        "desembolso": monto_desembolsado,
        "pct_desembolso": float(pct_desembolso),
        "productores": int(productores),
        "is_total": True,
    })

    # Departamentos con desembolso
    deptos_con_desembolso = datos.get("deptos_con_desembolso", 0)

    return {
        "fecha_corte": fecha_corte,
        "total_avisos": total_avisos,
        "avisos_lp": avisos_lp,
        "avisos_rimac": avisos_rimac,
        "top4_deptos": top4_deptos,
        "top4_total": top4_total,
        "top4_pct": top4_pct,
        "avisos_by_tipo": avisos_by_tipo,
        "top3_tipos": top3_tipos,
        "top3_total": top3_total,
        "top3_pct": top3_pct,
        "total_ajustados": total_ajustados,
        "pct_ajustados": pct_ajustados,
        "ajust_lp": ajust_lp,
        "ajust_rimac": ajust_rimac,
        "pct_ajust_lp": pct_ajust_lp,
        "pct_ajust_rimac": pct_ajust_rimac,
        "monto_indemnizado": monto_indemnizado,
        "idx_general": idx_general,
        "idx_lp": idx_lp,
        "idx_rimac": idx_rimac,
        "tabla_siniestralidad": tabla_siniestralidad,
        "tabla_coberturas": tabla_coberturas,
        "tabla_priorizados": tabla_priorizados,
        "tabla_desembolsos": tabla_desembolsos,
        "monto_desembolsado": monto_desembolsado,
        "productores": productores,
        "pct_desembolso": pct_desembolso,
        "deptos_con_desembolso": deptos_con_desembolso,
    }


# ═══════════════════════════════════════════════════════════════════
# GENERADOR DEL DOCUMENTO
# ═══════════════════════════════════════════════════════════════════

def generate_operatividad_docx(datos):
    """
    Genera el documento Word de Ayuda Memoria Operatividad SAC.

    Args:
        datos: dict de data_processor.process_dynamic_data()
    Returns:
        bytes: contenido del archivo .docx
    """
    d = _prepare_operatividad_data(datos)
    doc = Document()

    # Configurar márgenes
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # ═══════════════════════════════════════════════════════════════
    # TÍTULO
    # ═══════════════════════════════════════════════════════════════
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_after = Pt(2)
    run = t.add_run("AYUDA MEMORIA OPERATIVIDAD SAC")
    run.font.name = "Arial"
    run.font.size = Pt(13)
    run.font.bold = True
    run.underline = True

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2.paragraph_format.space_after = Pt(2)
    run = t2.add_run("CAMPAÑA AGRÍCOLA 2025-2026")
    run.font.name = "Arial"
    run.font.size = Pt(13)
    run.font.bold = True
    run.underline = True

    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t3.paragraph_format.space_after = Pt(12)
    run = t3.add_run(f"(AL {d['fecha_corte']})")
    run.font.name = "Arial"
    run.font.size = Pt(13)
    run.font.bold = True
    run.underline = True

    # ═══════════════════════════════════════════════════════════════
    # 1. SOBRE LA OPERATIVIDAD
    # ═══════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("1.  Sobre la operatividad de las empresas de seguros en la campaña 2025-2026:")
    run.font.name = "Arial"
    run.font.size = Pt(10)

    # ─── a) Avisos de siniestros ───
    _add_heading(doc, "a) Avisos de siniestros", level=2)

    _add_bullet(doc,
        f"Al {d['fecha_corte']}, se registran {fmt_int(d['total_avisos'])} avisos de siniestros "
        f"reportados por las DRA's/GRA's, que fueron recepcionados por las dos (2) empresas "
        f"de seguros de acuerdo al siguiente detalle:"
    )

    _add_subbullet(doc, f"La Positiva: {fmt_int(d['avisos_lp'])} avisos.")
    _add_subbullet(doc, f"Rímac: {fmt_int(d['avisos_rimac'])} avisos.")

    # Top 4 departamentos
    top4 = d["top4_deptos"]
    n_top = len(top4)
    top4_items = []
    for depto, count in top4.items():
        pct_of_top = (count / d["top4_total"] * 100) if d["top4_total"] > 0 else 0
        top4_items.append(f"{depto.title()} ({fmt_int(count)}, {pct_of_top:.1f}%)")

    _add_bullet(doc,
        f"El {d['top4_pct']:.1f}% de avisos reportados se concentra en {n_top} "
        f"departamentos: {fmt_int(d['top4_total'])} avisos distribuidos en "
        f"{', '.join(top4_items[:-1])} y {top4_items[-1]} avisos, respectivamente."
    )

    # Top 3 tipos de siniestro
    top3 = d["top3_tipos"]
    n_top3 = len(top3)
    top3_items = []
    for tipo, count in top3.items():
        pct_of_top3 = (count / d["top3_total"] * 100) if d["top3_total"] > 0 else 0
        top3_items.append(f"{tipo.lower()} ({fmt_int(count)}, {pct_of_top3:.1f}%)")

    _add_bullet(doc,
        f"El {d['top3_pct']:.1f}% de avisos reportados se concentra en {n_top3} siniestros: "
        f"{fmt_int(d['top3_total'])} avisos distribuidos en "
        f"{', '.join(top3_items[:-1])} y {top3_items[-1]} avisos."
    )

    # ─── b) Resultados ───
    _add_heading(doc, "b) Resultados", level=2)

    _add_bullet(doc,
        f"Del total de avisos de siniestros reportados ({fmt_int(d['total_avisos'])}) a las empresas "
        f"de seguros, al {d['fecha_corte']} se han ajustado y evaluado {fmt_int(d['total_ajustados'])} "
        f"avisos que representa el {d['pct_ajustados']:.2f}% de los avisos."
    )

    _add_bullet(doc, "Los ajustes y evaluación de los mismos por cada empresa de seguros es la siguiente:")

    _add_subbullet(doc,
        f"La Positiva: {fmt_int(d['ajust_lp'])} ajustes de {fmt_int(d['avisos_lp'])} "
        f"avisos de siniestros, {d['pct_ajust_lp']:.2f}% de avisos atendidos."
    )
    _add_subbullet(doc,
        f"Rímac: {fmt_int(d['ajust_rimac'])} ajustes de {fmt_int(d['avisos_rimac'])} "
        f"avisos de siniestros, {d['pct_ajust_rimac']:.2f}% de avisos atendidos."
    )

    # Indemnizaciones
    p = _add_bullet(doc, "")
    p.clear()
    run = p.add_run(f"Las indemnizaciones reconocidas a la fecha, por parte de las empresas de seguros, equivalen al monto de {fmt(d['monto_indemnizado'])}.")
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.bold = True
    run.underline = True

    # Índice de siniestralidad
    _add_bullet(doc,
        f"El índice de siniestralidad (costo de los siniestros entre la prima neta, sin incluir IGV, "
        f"de la póliza) en lo que va de la campaña, es de {d['idx_general']:.2f}%, siendo el índice "
        f"para cada empresa de seguros el siguiente:"
    )
    _add_subbullet(doc, f"La Positiva Seguros: {d['idx_lp']:.2f}%.")
    _add_subbullet(doc, f"Rímac Seguros: {d['idx_rimac']:.2f}%.")

    _add_body(doc,
        "Es preciso resaltar que la vigencia de la póliza es desde el 01 de agosto de 2025 "
        "hasta 01 de agosto de 2026, por lo que, todos los eventos adversos que se presenten "
        "durante este periodo, que afecten los cultivos asegurados, sean reportados por las "
        "DRAs/GRAs a las empresas de seguros, para que estas realicen las evaluaciones "
        "correspondientes.",
        italic=True, size=9,
    )

    _add_body(doc, "El índice de siniestralidad para cada departamento se detalla en el siguiente cuadro:", size=10)

    # ═══════════════════════════════════════════════════════════════
    # TABLA: SINIESTRALIDAD POR DEPARTAMENTO
    # ═══════════════════════════════════════════════════════════════
    headers_sin = ["EMPRESA DE SEGUROS", "DEPARTAMENTO", "INDEMNIZACIÓN (S/)",
                   "SUPERFICIE INDEMNIZADA (Has)", "PRIMA NETA (S/)", "ÍNDICE SINIESTRALIDAD (%)"]
    col_widths_sin = [2000, 1700, 1700, 1700, 1700, 1500]

    n_rows = len(d["tabla_siniestralidad"]) + 1  # +1 header
    table = doc.add_table(rows=n_rows, cols=6)
    table.autofit = False

    # Header
    for i, h in enumerate(headers_sin):
        cell = table.rows[0].cells[i]
        _write_cell(cell, h, bold=True, size=7, align=WD_ALIGN_PARAGRAPH.CENTER,
                    font_color=C["WHITE"], bg_color=C["HEADER_BG"])
        _set_cell_width(cell, col_widths_sin[i])

    # Data rows
    for row_idx, row_data in enumerate(d["tabla_siniestralidad"]):
        is_sub = row_data.get("is_subtotal", False)
        is_tot = row_data.get("is_total", False)
        bg = C["HEADER_BG"] if is_tot else (C["SUBTOTAL_BG"] if is_sub else (C["ALT_ROW"] if row_idx % 2 == 0 else None))
        fc = C["WHITE"] if is_tot else C["BLACK"]
        b = is_sub or is_tot

        r = table.rows[row_idx + 1]
        _write_cell(r.cells[0], row_data["empresa"], bold=b, size=7, bg_color=bg, font_color=fc)
        _write_cell(r.cells[1], row_data["departamento"], bold=b, size=7, bg_color=bg, font_color=fc)
        _write_cell(r.cells[2], fmt(row_data["indemnizacion"]), bold=b, size=7,
                    align=WD_ALIGN_PARAGRAPH.RIGHT, bg_color=bg, font_color=fc)
        _write_cell(r.cells[3], fmt_n(row_data["sup_indemnizada"]), bold=b, size=7,
                    align=WD_ALIGN_PARAGRAPH.RIGHT, bg_color=bg, font_color=fc)
        _write_cell(r.cells[4], fmt(row_data["prima_neta"]), bold=b, size=7,
                    align=WD_ALIGN_PARAGRAPH.RIGHT, bg_color=bg, font_color=fc)
        _write_cell(r.cells[5], fmt_pct(row_data["indice"]), bold=b, size=7,
                    align=WD_ALIGN_PARAGRAPH.CENTER, bg_color=bg, font_color=fc)

    doc.add_paragraph()  # spacer

    # ═══════════════════════════════════════════════════════════════
    # TABLA: COBERTURAS
    # ═══════════════════════════════════════════════════════════════
    if d["tabla_coberturas"]:
        total_comp = sum(r.get("complementaria", 0) for r in d["tabla_coberturas"] if not r.get("is_subtotal") and not r.get("is_total"))
        total_cat = sum(r.get("catastrofica", 0) for r in d["tabla_coberturas"] if not r.get("is_subtotal") and not r.get("is_total"))
        pct_comp = (total_comp / d["monto_indemnizado"] * 100) if d["monto_indemnizado"] > 0 else 0
        pct_cat = (total_cat / d["monto_indemnizado"] * 100) if d["monto_indemnizado"] > 0 else 0

        _add_bullet(doc,
            f"Del total de las indemnizaciones reconocidas por las empresas de seguros, "
            f"se tiene que el {pct_comp:.0f}% es por la cobertura complementaria "
            f"({fmt(total_comp)}); el {pct_cat:.0f}% restante es por la cobertura "
            f"catastrófica de evaluación ({fmt(total_cat)})."
        )

        headers_cob = ["EMPRESA DE SEGUROS", "DEPARTAMENTO", "COBERTURA COMPLEMENTARIA",
                       "COBERTURA CATASTRÓFICA", "INDEMNIZACIÓN TOTAL"]
        col_widths_cob = [1900, 1600, 2000, 2000, 1800]

        n_rows_c = len(d["tabla_coberturas"]) + 1
        table_c = doc.add_table(rows=n_rows_c, cols=5)
        table_c.autofit = False

        for i, h in enumerate(headers_cob):
            cell = table_c.rows[0].cells[i]
            _write_cell(cell, h, bold=True, size=7, align=WD_ALIGN_PARAGRAPH.CENTER,
                        font_color=C["WHITE"], bg_color=C["HEADER_BG"])
            _set_cell_width(cell, col_widths_cob[i])

        for row_idx, row_data in enumerate(d["tabla_coberturas"]):
            is_sub = row_data.get("is_subtotal", False)
            is_tot = row_data.get("is_total", False)
            bg = C["HEADER_BG"] if is_tot else (C["SUBTOTAL_BG"] if is_sub else (C["ALT_ROW"] if row_idx % 2 == 0 else None))
            fc = C["WHITE"] if is_tot else C["BLACK"]
            b = is_sub or is_tot

            r = table_c.rows[row_idx + 1]
            _write_cell(r.cells[0], row_data["empresa"], bold=b, size=7, bg_color=bg, font_color=fc)
            _write_cell(r.cells[1], row_data.get("departamento", ""), bold=b, size=7, bg_color=bg, font_color=fc)
            _write_cell(r.cells[2], fmt(row_data["complementaria"]), bold=b, size=7,
                        align=WD_ALIGN_PARAGRAPH.RIGHT, bg_color=bg, font_color=fc)
            _write_cell(r.cells[3], fmt(row_data["catastrofica"]), bold=b, size=7,
                        align=WD_ALIGN_PARAGRAPH.RIGHT, bg_color=bg, font_color=fc)
            _write_cell(r.cells[4], fmt(row_data["total"]), bold=b, size=7,
                        align=WD_ALIGN_PARAGRAPH.RIGHT, bg_color=bg, font_color=fc)

        doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════
    # CUADRO 4: CULTIVOS PRIORIZADOS
    # ═══════════════════════════════════════════════════════════════
    if d["tabla_priorizados"]:
        _add_body(doc, "Cuadro N° 04: Indemnizaciones por cultivos Priorizados y No Priorizados, SAC 2025-2026",
                  bold=True, size=10)

        headers_p = ["EMPRESA DE SEGUROS", "CULTIVOS", "SUPERFICIE INDEMNIZADA (Has)", "INDEMNIZACIÓN (S/)"]
        col_widths_p = [2200, 1800, 2600, 2700]

        n_rows_p = len(d["tabla_priorizados"]) + 1
        table_p = doc.add_table(rows=n_rows_p, cols=4)
        table_p.autofit = False

        for i, h in enumerate(headers_p):
            cell = table_p.rows[0].cells[i]
            _write_cell(cell, h, bold=True, size=7, align=WD_ALIGN_PARAGRAPH.CENTER,
                        font_color=C["WHITE"], bg_color=C["HEADER_BG"])
            _set_cell_width(cell, col_widths_p[i])

        for row_idx, row_data in enumerate(d["tabla_priorizados"]):
            is_sub = row_data.get("is_subtotal", False)
            is_tot = row_data.get("is_total", False)
            bg = C["HEADER_BG"] if is_tot else (C["SUBTOTAL_BG"] if is_sub else (C["ALT_ROW"] if row_idx % 2 == 0 else None))
            fc = C["WHITE"] if is_tot else C["BLACK"]
            b = is_sub or is_tot

            r = table_p.rows[row_idx + 1]
            _write_cell(r.cells[0], row_data["empresa"], bold=b, size=7, bg_color=bg, font_color=fc)
            _write_cell(r.cells[1], row_data.get("cultivo", ""), bold=b, size=7, bg_color=bg, font_color=fc)
            _write_cell(r.cells[2], fmt_n(row_data["sup_indemnizada"]), bold=b, size=7,
                        align=WD_ALIGN_PARAGRAPH.RIGHT, bg_color=bg, font_color=fc)
            _write_cell(r.cells[3], fmt(row_data["indemnizacion"]), bold=b, size=7,
                        align=WD_ALIGN_PARAGRAPH.RIGHT, bg_color=bg, font_color=fc)

        doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════
    # CUADRO 5: DESEMBOLSOS
    # ═══════════════════════════════════════════════════════════════
    _add_bullet(doc,
        f"Con respecto a las indemnizaciones reconocidas por parte de las empresas de seguros "
        f"del SAC 2025-2026, a la fecha el avance de desembolsos realizados es la siguiente:"
    )

    # Listar departamentos con desembolso
    deptos_lp_desemb = [r["departamento"] for r in d["tabla_desembolsos"]
                        if r.get("empresa") == "LA POSITIVA" and r.get("desembolso", 0) > 0
                        and not r.get("is_subtotal") and not r.get("is_total")]
    deptos_rimac_desemb = [r["departamento"] for r in d["tabla_desembolsos"]
                           if r.get("empresa") == "RÍMAC" and r.get("desembolso", 0) > 0
                           and not r.get("is_subtotal") and not r.get("is_total")]

    if deptos_lp_desemb:
        _add_body(doc,
            f"A nivel de departamento, La Positiva seguros ha iniciado con los desembolsos de "
            f"las indemnizaciones en los departamentos de {', '.join(deptos_lp_desemb)}.",
            size=10,
        )
    if deptos_rimac_desemb:
        _add_body(doc,
            f"Rímac ha iniciado con los desembolsos de indemnizaciones en los departamentos de "
            f"{', '.join(deptos_rimac_desemb)}.",
            size=10,
        )

    _add_body(doc, "Cuadro N° 05: Desembolsos y número de productores, SAC 2025-2026",
              bold=True, size=10)

    headers_d = ["EMPRESA DE SEGUROS", "DEPARTAMENTO", "INDEMNIZACIÓN (S/)",
                 "DESEMBOLSO (S/)", "% DESEMBOLSO", "N° PRODUCTORES CON DESEMBOLSO"]
    col_widths_d = [1800, 1500, 1700, 1700, 1200, 1400]

    n_rows_d = len(d["tabla_desembolsos"]) + 1
    table_d = doc.add_table(rows=n_rows_d, cols=6)
    table_d.autofit = False

    for i, h in enumerate(headers_d):
        cell = table_d.rows[0].cells[i]
        _write_cell(cell, h, bold=True, size=7, align=WD_ALIGN_PARAGRAPH.CENTER,
                    font_color=C["WHITE"], bg_color=C["HEADER_BG"])
        _set_cell_width(cell, col_widths_d[i])

    for row_idx, row_data in enumerate(d["tabla_desembolsos"]):
        is_sub = row_data.get("is_subtotal", False)
        is_tot = row_data.get("is_total", False)
        bg = C["HEADER_BG"] if is_tot else (C["SUBTOTAL_BG"] if is_sub else (C["ALT_ROW"] if row_idx % 2 == 0 else None))
        fc = C["WHITE"] if is_tot else C["BLACK"]
        b = is_sub or is_tot

        r = table_d.rows[row_idx + 1]
        _write_cell(r.cells[0], row_data["empresa"], bold=b, size=7, bg_color=bg, font_color=fc)
        _write_cell(r.cells[1], row_data.get("departamento", ""), bold=b, size=7, bg_color=bg, font_color=fc)
        _write_cell(r.cells[2], fmt(row_data["indemnizacion"]), bold=b, size=7,
                    align=WD_ALIGN_PARAGRAPH.RIGHT, bg_color=bg, font_color=fc)
        _write_cell(r.cells[3], fmt(row_data["desembolso"]), bold=b, size=7,
                    align=WD_ALIGN_PARAGRAPH.RIGHT, bg_color=bg, font_color=fc)
        _write_cell(r.cells[4], fmt_pct(row_data["pct_desembolso"]), bold=b, size=7,
                    align=WD_ALIGN_PARAGRAPH.CENTER, bg_color=bg, font_color=fc)
        _write_cell(r.cells[5], fmt_int(row_data["productores"]) if row_data["productores"] > 0 else "-",
                    bold=b, size=7, align=WD_ALIGN_PARAGRAPH.CENTER, bg_color=bg, font_color=fc)

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════
    # PÁRRAFOS FINALES
    # ═══════════════════════════════════════════════════════════════
    _add_bullet(doc,
        f"A la fecha se van indemnizando a {fmt_int(d['productores'])} productores en "
        f"{d['deptos_con_desembolso']} de los 24 departamentos por el Seguro Agrícola "
        f"Catastrófico en la presente campaña."
    )

    _add_bullet(doc,
        f"Se va teniendo un porcentaje de desembolso del {float(d['pct_desembolso']):.2f}%."
    )

    # ═══ GUARDAR ═══
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()
